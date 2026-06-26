from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ====== 全局样式设置 ======
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ====== 辅助函数 ======
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(50, 50, 50)
    return p

# ====== 标题 ======
title = doc.add_heading('Java科学计算器实验报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)

# 基本信息
info_items = [
    ('课程名称', 'Java面向对象程序设计'),
    ('实验项目', '科学计算器综合设计'),
    ('实验类型', '综合性设计实验'),
]
for label, value in info_items:
    p = doc.add_paragraph()
    run1 = p.add_run(f'{label}：')
    run1.bold = True
    run2 = p.add_run(value)

doc.add_page_break()

# ====== 一、实验目的 ======
add_heading_styled('一、实验目的', level=1)
purposes = [
    '掌握Java面向对象编程的核心概念：类与对象、继承与多态，能够灵活运用抽象类、接口、方法重写等机制进行程序开发。',
    '掌握数组与集合（如ArrayList）的使用、异常处理机制（try-catch-finally）、图形用户界面（Swing/JFrame）、文件I/O（BufferedReader/BufferedWriter）及常用工具类的应用。',
    '能够综合运用所学知识，独立设计并实现一个完整的科学计算器应用程序，涵盖控制台版本和图形界面版本。',
]
for i, p_text in enumerate(purposes, 1):
    add_para(f'{i}. {p_text}')

# ====== 二、实验内容 ======
add_heading_styled('二、实验内容', level=1)

add_heading_styled('实验一：面向对象的计算器核心设计', level=2)
add_para(
    '本实验以面向对象编程思想为指导，设计了计算器的核心运算框架。'
    '首先定义了一个抽象的Operation基类，封装了二元运算的公共接口calculate(double, double)和getSymbol()方法，'
    '同时提供了makeRecord方法用于生成运算记录字符串。',
    indent=True
)
add_para(
    '基于Operation基类，通过继承实现了以下具体运算类：'
    'Add（加法）、Sub（减法）、Mul（乘法）、Div（除法）、Mod（取余）、Pow（乘方）等二元运算类，'
    '以及Sin（正弦）、Cos（余弦）、Tan（正切）、Sqrt（平方根）等一元运算类。'
    '一元运算类通过重写makeRecord方法，实现了与二元运算不同的记录格式。',
    indent=True
)
add_para(
    '在运算调度方面，设计了Calculator类作为运算执行器，CalculatorEngine类作为计算引擎，'
    '负责管理运算状态（当前输入值、当前运算符、是否为新输入等），实现了运算流程的完整控制。'
    '同时设计了HistoryManager类，使用ArrayList<String>集合存储运算历史记录，实现了历史记录的添加、查看和导出功能。',
    indent=True
)

add_heading_styled('实验二：GUI科学计算器', level=2)
add_para(
    '本实验运用Swing图形界面库，构建了可视化科学计算器应用程序。'
    '主窗口CalculatorGUI继承自JFrame，采用BorderLayout布局管理器，将界面划分为显示区（北区）、按钮区（中间）和功能区（南区）三个部分。',
    indent=True
)
add_para(
    '显示区使用JTextField组件，设置右对齐和大号字体，用于显示输入和计算结果。'
    '按钮区使用6×5的GridLayout网格布局，包含数字按钮（0-9）、运算符按钮（+−×÷%^）、'
    '科学函数按钮（sin/cos/tan/√）、控制按钮（C/del/=）和辅助按钮（括号、小数点）。'
    '功能区提供了"显示历史"和"保存/加载"两个功能按钮。',
    indent=True
)
add_para(
    '事件处理采用匿名内部类实现ActionListener接口，为每个按钮注册独立的事件处理器。'
    '计算结果的显示使用FormatUtil工具类进行格式化，保留6位小数并去除多余的零。'
    '历史记录支持保存到文本文件（FileWriter）和从文件加载（FileReader），实现了数据的持久化存储。',
    indent=True
)
add_para(
    '同时，程序还包含一个控制台版本的Main类，使用Scanner读取用户输入，'
    '支持交互式的命令行计算体验，展示了同一套运算核心在不同界面下的复用能力。',
    indent=True
)

# ====== 三、核心类设计与关键代码 ======
add_heading_styled('三、核心类设计与关键代码', level=1)

# 3.1 类图
add_heading_styled('3.1 类图', level=2)
add_para('以下是科学计算器的核心类图设计：')
add_code(
    '┌─────────────────────────────────────────────────────┐\n'
    '│                   类图 (Class Diagram)                │\n'
    '├─────────────────────────────────────────────────────┤\n'
    '│                                                     │\n'
    '│  ┌──────────────────────┐                           │\n'
    '│  │  <<abstract>>        │                           │\n'
    '│  │     Operation        │                           │\n'
    '│  ├──────────────────────┤  ← 继承                    │\n'
    '│  │ + calculate(a,b):double│  ┌──────┬──────┬──────┐ │\n'
    '│  │ + getSymbol():String │  │ Add  │ Sub  │ Mul  │ │\n'
    '│  │ + makeRecord(...)    │  ├──────┼──────┼──────┤ │\n'
    '│  └──────────────────────┘  │ Div  │ Mod  │ Pow  │ │\n'
    '│         ↑                  ├──────┼──────┼──────┤ │\n'
    '│         │ 继承             │ Sin  │ Cos  │ Tan  │ │\n'
    '│         │                  │ Sqrt │      │      │ │\n'
    '│         │                  └──────┴──────┴──────┘ │\n'
    '│  ┌──────────────────────┐                           │\n'
    '│  │     Calculator       │  ──→ 使用                 │\n'
    '│  ├──────────────────────┤                           │\n'
    '│  │ + compute(op,a,b)    │         ┌──────────────┐  │\n'
    '│  │ - result: double     │         │CalculatorEngine│ │\n'
    '│  └──────────────────────┘    ──→ ├──────────────┤  │\n'
    '│                                    │ + setOperator()│  │\n'
    '│  ┌──────────────────────┐         │ + calculate()  │  │\n'
    '│  │   FormatUtil         │         │ + executeUnary()│ │\n'
    '│  ├──────────────────────┤         └──────────────┘  │\n'
    '│  │ + formatResult():Str │                           │\n'
    '│  └──────────────────────┘         ┌──────────────┐  │\n'
    '│                                    │HistoryManager│  │\n'
    '│  ┌──────────────────────┐         ├──────────────┤  │\n'
    '│  │   CalculatorGUI      │──→ 使用  │+ addRecord()  │  │\n'
    '│  ├──────────────────────┤         │+ getAllRecords│  │\n'
    '│  │  (Swing JFrame)      │         │+ toPlainString│  │\n'
    '│  │  + displayField      │         └──────────────┘  │\n'
    '│  │  + buttonPanel       │                            │\n'
    '│  │  + historyManager    │                            │\n'
    '│  └──────────────────────┘                            │\n'
    '└─────────────────────────────────────────────────────┘'
)

# 3.2 关键代码片段
add_heading_styled('3.2 关键代码片段', level=2)

# 片段1：Operation抽象类
add_para('【代码片段1】Operation抽象基类 —— 体现抽象与封装', bold=True)
add_para(
    'Operation是运算体系的基类，定义了所有运算类的公共接口。'
    'calculate方法采用抽象方法设计，强制子类实现具体的运算逻辑，体现了面向对象中的抽象原则。',
    indent=True
)

code1 = '''public abstract class Operation {
    public abstract double calculate(double a, double b);
    public abstract String getSymbol();

    // 生成计算记录（二元运算）
    public String makeRecord(double a, double b, double result) {
        return FormatUtil.formatResult(a) + " " + getSymbol() + " "
                + FormatUtil.formatResult(b) + " = " + FormatUtil.formatResult(result);
    }
}'''
add_code(code1)

# 片段2：继承与多态（Add / Sin示例）
add_para('【代码片段2】继承与多态 —— Add类与Sin类', bold=True)
add_para(
    'Add和Sin都继承自Operation，体现了继承机制。Add的calculate实现两数相加；'
    'Sin作为一元运算，只使用第一个参数a，并通过重写makeRecord方法改变记录格式，体现了多态性。',
    indent=True
)

code2 = '''// Add.java —— 二元运算
public class Add extends Operation {
    @Override
    public double calculate(double a, double b) {
        return a + b;
    }
    @Override
    public String getSymbol() { return "+"; }
}

// Sin.java —— 一元运算（多态：重写makeRecord）
public class Sin extends Operation {
    @Override
    public double calculate(double a, double b) {
        return Math.sin(a);  // 忽略参数b，只使用a
    }
    @Override
    public String getSymbol() { return "sin"; }
    @Override
    public String makeRecord(double a, double b, double result) {
        return "sin(" + FormatUtil.formatResult(a) + ") = "
               + FormatUtil.formatResult(result);
    }
}'''
add_code(code2)

# 片段3：异常处理
add_para('【代码片段3】异常处理 —— Div类', bold=True)
add_para(
    '在除法运算中，通过ArithmeticException异常来处理除数为零的情况。'
    '当除数为0时，抛出包含错误信息的异常，由上层GUI捕获并弹出错误提示框。'
    '同样，Sqrt类中也对负数开平方进行了异常处理。',
    indent=True
)

code3 = '''// Div.java —— 除法运算，含除零检查
public class Div extends Operation {
    @Override
    public double calculate(double a, double b) {
        if(b == 0){
            throw new ArithmeticException("除数不能为0");
        }
        return a / b;
    }
    @Override
    public String getSymbol() { return "/"; }
}'''
add_code(code3)

# 片段4：计算引擎
add_para('【代码片段4】CalculatorEngine计算引擎', bold=True)
add_para(
    'CalculatorEngine是运算的核心调度器，管理运算状态。'
    '它包含当前操作数(num1)、当前运算符(currentOp)和新输入标志(isNewInput)。'
    'calculate方法执行二元运算并返回CalcResult对象（含计算结果和记录字符串），'
    'executeUnary方法执行一元运算。CalcResult作为内部类封装了运算结果。',
    indent=True
)

code4 = '''public class CalculatorEngine {
    private Calculator calculator;
    private double num1;
    private Operation currentOp;
    private boolean isNewInput;

    // 设置运算操作符，保存当前值
    public void setOperator(Operation op, double currentValue) {
        num1 = currentValue;
        currentOp = op;
        isNewInput = true;
    }

    // 执行二元运算
    public CalcResult calculate(double num2) throws ArithmeticException {
        if (currentOp == null) {
            throw new IllegalStateException("没有选择运算符");
        }
        double result = calculator.compute(currentOp, num1, num2);
        String record = currentOp.makeRecord(num1, num2, result);
        this.num1 = result;
        this.currentOp = null;
        this.isNewInput = true;
        return new CalcResult(result, record);
    }

    // 执行一元运算（开方、sin、cos、tan）
    public CalcResult executeUnary(Operation op, double value) throws ArithmeticException {
        double result = op.calculate(value, 0);
        String record = op.makeRecord(value, 0, result);
        this.isNewInput = true;
        return new CalcResult(result, record);
    }

    // 内部类：封装计算结果
    public static class CalcResult {
        private final double value;
        private final String record;
        // ...
    }
}'''
add_code(code4)

# 片段5：Swing GUI
add_para('【代码片段5】Swing图形用户界面 —— CalculatorGUI', bold=True)
add_para(
    'CalculatorGUI继承自JFrame，使用BorderLayout布局。'
    '显示区使用JTextField（36号字体、右对齐），按钮区使用6×5 GridLayout网格布局。'
    '通过匿名内部类实现ActionListener接口为按钮注册事件。'
    'GUI界面统一使用了"微软雅黑"字体，保证界面美观。',
    indent=True
)

code5 = '''public class CalculatorGUI extends JFrame {
    private JTextField displayField;
    private JButton[] numberButtons;
    private CalculatorEngine engine;
    private HistoryManager historyManager;

    public CalculatorGUI() {
        engine = new CalculatorEngine();
        historyManager = new HistoryManager();

        setTitle("Java GUI 计算器");
        setDefaultCloseOperation(JFrame.EXIT_ON_CLOSE);
        setSize(400, 600);
        setLocationRelativeTo(null);
        setLayout(new BorderLayout(10, 10));

        // 显示面板
        displayField = new JTextField();
        displayField.setFont(new Font("微软雅黑", Font.BOLD, 36));
        displayField.setHorizontalAlignment(JTextField.RIGHT);
        displayField.setEditable(false);
        displayField.setText("0");

        // 按钮面板 (6行5列)
        JPanel buttonPanel = new JPanel(new GridLayout(6, 5, 5, 5));
        // ... 添加各种按钮 ...

        // 事件处理
        for (int i = 0; i < 10; i++) {
            final int digit = i;
            numberButtons[i].addActionListener(e -> appendDigit(digit));
        }
        equalsBtn.addActionListener(e -> calculateResult());
        // ...
    }
    // 使用 SwingUtilities 确保线程安全
    public static void main(String[] args) {
        SwingUtilities.invokeLater(() -> new CalculatorGUI());
    }
}'''
add_code(code5)

# 片段6：集合与文件I/O
add_para('【代码片段6】ArrayList集合与文件I/O —— HistoryManager', bold=True)
add_para(
    'HistoryManager使用ArrayList<String>集合存储运算历史记录，支持动态扩容。'
    '通过FileWriter/BufferedWriter将历史记录写入文本文件，通过FileReader/BufferedReader从文件读取，'
    '实现了运算历史的持久化存储与恢复功能，体现了文件I/O的应用。',
    indent=True
)

code6 = '''public class HistoryManager {
    private ArrayList<String> records;  // 使用ArrayList存储历史

    public void addRecord(String record) {
        records.add(record);
    }

    public String toPlainString() {
        StringBuilder sb = new StringBuilder();
        for (int i = 0; i < records.size(); i++) {
            sb.append(i + 1).append(". ").append(records.get(i)).append("\\n");
        }
        return sb.toString();
    }
}

// 保存历史到文件 (FileWriter + BufferedWriter)
try (BufferedWriter writer = new BufferedWriter(new FileWriter("history.txt"))) {
    writer.write("===== 计算历史 =====\\n");
    for (String record : historyManager.getAllRecords()) {
        writer.write(record + "\\n");
    }
} catch (IOException e) {
    JOptionPane.showMessageDialog(this, "保存失败: " + e.getMessage(),
                                   "错误", JOptionPane.ERROR_MESSAGE);
}

// 从文件加载历史 (FileReader + BufferedReader)
try (BufferedReader reader = new BufferedReader(new FileReader("history.txt"))) {
    String line;
    while ((line = reader.readLine()) != null) {
        content.append(line).append("\\n");
    }
} catch (IOException e) {
    JOptionPane.showMessageDialog(this, "读取失败", "错误", JOptionPane.ERROR_MESSAGE);
}'''
add_code(code6)

# 片段7：FormatUtil
add_para('【代码片段7】FormatUtil工具类 —— 常用类的应用', bold=True)
add_para(
    'FormatUtil工具类负责格式化计算结果，当结果为整数时显示为整数格式，'
    '为小数时保留6位有效数字并去除末尾多余的零。使用了String.format和正则表达式。',
    indent=True
)

code7 = '''public class FormatUtil {
    public static String formatResult(double value) {
        if (value == (long) value) {
            return String.valueOf((long) value);
        }
        return String.format("%.6f", value)
               .replaceAll("0*$", "").replaceAll("\\\\.$", "");
    }
}'''
add_code(code7)

doc.add_page_break()

# ====== 四、实验结果 ======
add_heading_styled('四、实验结果', level=1)
add_para('在IntelliJ IDEA开发环境中运行程序，得到以下实验结果：')
add_para('（请在以下位置插入程序运行截图）', bold=True)

results = [
    ('图1：程序启动界面', '科学计算器启动后的初始界面，显示数字按钮、运算符按钮、科学函数按钮（sin/cos/tan/√）以及历史功能按钮。'),
    ('图2：四则运算演示', '进行加法运算（如 123 + 456 = 579）的界面截图。'),
    ('图3：科学函数运算演示', '进行三角函数运算（如 sin(90) = 1）和开方运算（如 √(16) = 4）的界面截图。'),
    ('图4：除零异常处理', '当除数为0时，程序弹出错误提示框"除数不能为0"的截图。'),
    ('图5：历史记录功能', '点击"显示历史"按钮后，弹出历史记录对话框的截图。'),
    ('图6：文件保存与加载', '保存历史记录到history.txt文件，以及从文件加载历史记录的截图。'),
]
for fig_name, fig_desc in results:
    add_para(f'  {fig_name}：{fig_desc}')

# ====== 五、遇到的问题与解决方法 ======
add_heading_styled('五、遇到的问题与解决方法', level=1)

problems = [
    (
        '问题1：一元运算的calculate方法设计',
        '最初设计Operation的calculate方法时，所有运算都接受两个参数。但对于sin、cos等一元运算，第二个参数是多余的。'
        '如果直接舍弃第二个参数又不使用，代码不够优雅。',
        '在Sin、Cos、Tan和Sqrt类中，虽然calculate方法签名仍然接受两个参数，但重写了makeRecord方法，'
        '使一元运算的记录格式与二元运算不同。同时，在CalculatorEngine中专门设计了executeUnary方法来处理一元运算，'
        '与二元运算的calculate方法分开调用，从调用层面保证了一元运算的正确性。'
    ),
    (
        '问题2：界面中文乱码',
        '在Swing界面中使用中文标签时，可能会出现显示为方框或乱码的情况。',
        '统一将界面字体设置为"微软雅黑"（Microsoft YaHei），该字体对中文支持良好。'
        '同时使用setFont方法为所有按钮和文本框指定相同的字体，保证了界面语言的正确显示和风格统一。'
    ),
    (
        '问题3：计算结果格式化',
        '科学计算涉及浮点数运算，直接显示原始double值会导致结果过长（如3.141592653589793）或出现科学计数法。'
        '同时，整数运算如2+3=5时也不希望显示为5.000000。',
        '设计FormatUtil工具类，先判断值是否为整数（value == (long)value），'
        '整数直接显示为整数格式；小数则使用String.format保留6位小数，'
        '并通过正则表达式replaceAll("0*$", "").replaceAll("\\\\.$", "")去除末尾多余的零和小数点。'
    ),
    (
        '问题4：历史记录的持久化存储',
        '程序关闭后，内存中的历史记录会丢失。需要一种机制将历史记录保存到磁盘。',
        '使用FileWriter和BufferedWriter将历史记录写入history.txt文本文件，'
        '使用FileReader和BufferedReader从文件读取历史记录。采用try-with-resources语法确保资源自动关闭。'
        '同时在GUI中通过"保存/加载"按钮提供交互界面，使用JOptionPane的选项对话框让用户选择保存或加载。'
    ),
    (
        '问题5：布局空间不足',
        '计算器功能较多（包括科学函数按钮），6×5的GridLayout需要合理安排按钮位置，'
        '不能显得拥挤或杂乱。',
        '经过多次调整，最终确定的布局为：第一行为科学函数行（sin/cos/tan/(/)），'
        '第二行开始为数字和运算符区域，并使用空的JPanel作为占位符填充保留位置，'
        '使界面布局均匀有序。'
    ),
]

for i, (title_text, problem_desc, solution) in enumerate(problems, 1):
    add_para(f'{title_text}', bold=True)
    add_para(f'问题描述：{problem_desc}', indent=True)
    add_para(f'解决方法：{solution}', indent=True)

# ====== 六、实验总结与反思 ======
add_heading_styled('六、实验总结与反思', level=1)
add_para(
    '通过本次科学计算器的设计与实现，我全面巩固了Java面向对象编程的核心知识，'
    '包括类的设计、继承与多态、抽象类的应用等。Operation抽象类的设计使运算逻辑具有良好的可扩展性，'
    '新增运算类型只需继承Operation并实现calculate方法，无需修改现有代码，体现了开闭原则。',
    indent=True
)
add_para(
    '在Swing图形界面方面，我掌握了JFrame、JPanel、JTextField、JButton等常用组件的使用，'
    '了解了BorderLayout、GridLayout等布局管理器的特点，以及事件驱动编程模型（ActionListener接口）的应用。'
    '通过匿名内部类和Lambda表达式两种方式实现了按钮事件的绑定。',
    indent=True
)
add_para(
    '在异常处理方面，我在Div和Sqrt等运算类中加入了输入合法性检查，'
    '在GUI层通过try-catch捕获异常并使用JOptionPane显示友好的错误提示信息，'
    '提升了程序的健壮性和用户体验。',
    indent=True
)
add_para(
    '在集合与文件I/O方面，我使用ArrayList<String>管理运算历史记录，'
    '利用FileWriter/BufferedWriter将历史保存到文件，利用FileReader/BufferedReader从文件加载，'
    '实现了数据的持久化存储。',
    indent=True
)
add_para(
    '通过本次实验，我也认识到程序仍有一些可以改进的地方。'
    '例如，目前括号功能仅支持手动输入显示，尚未实现完整的表达式解析求值；'
    '科学函数的输入目前基于角度制，可以添加角度/弧度切换功能；'
    '界面配色和按钮样式较为基础，可以进一步美化为更现代化的设计风格。'
    '此外，使用java.math.BigDecimal替代double类型进行计算，可以避免浮点数精度损失的问题。'
    '这些将在后续的学习中不断完善。',
    indent=True
)
add_para(
    '总的来说，本次实验不仅加深了我对Java面向对象编程的理解，'
    '也锻炼了我独立设计、编码和调试程序的能力，为后续更复杂的软件开发打下了坚实基础。',
    indent=True
)

# 保存
output_path = 'C:\\Users\\33284\\AppData\\Local\\Claude-3p\\local-agent-mode-sessions\\ecf8421b-092a-4d79-abab-8a9a864e137d\\00000000-0000-4000-8000-000000000001\\local_f7244120-4956-414a-8b1e-90cbebb9077a\\outputs\\Java科学计算器实验报告.docx'
doc.save(output_path)
print(f"报告已生成: {output_path}")
